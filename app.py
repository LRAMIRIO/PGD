import pandas as pd
from docx import Document
import nltk
import streamlit as st
from io import BytesIO

# Baixa tokenizer do NLTK
nltk.download('punkt')
nltk.download('punkt_tab')

# Função para criar texto formal
def gerar_texto_formal(lista_descricoes):
    lista_limpa = [str(x).strip() for x in lista_descricoes if pd.notna(x) and str(x).strip() != ""]
    if not lista_limpa:
        return ""
    texto_unido = " ".join(lista_limpa)
    texto_unido = texto_unido.replace("  ", " ").strip()
    frases = nltk.sent_tokenize(texto_unido, language='portuguese')
    if len(frases) > 1:
        texto_formal = frases[0]
        for frase in frases[1:]:
            texto_formal += " Além disso, " + frase[0].lower() + frase[1:]
    else:
        texto_formal = texto_unido
    return texto_formal[0].upper() + texto_formal[1:]

# Interface Streamlit
st.title("📄 Gerador de Relatórios Formais")
st.write("Envie uma planilha Excel com qualquer número de colunas (temas na primeira linha e descrições nas linhas seguintes).")

arquivo = st.file_uploader("📂 Enviar planilha (.xlsx)", type=["xlsx"])

if arquivo is not None:
    df = pd.read_excel(arquivo, header=None)

    relatorios = {}
    for col in df.columns:
        tema = str(df.iloc[0, col]).strip()
        if tema == "" or tema.lower() == "nan":
            continue
        conteudo = df.iloc[1:, col].tolist()
        texto_formal = gerar_texto_formal(conteudo)
        if texto_formal:
            relatorios[tema] = texto_formal

    if relatorios:
        st.subheader("📄 Relatório Gerado")
        for tema, texto in relatorios.items():
            st.markdown(f"### {tema}")
            st.write(texto)

        # Salva em Word
        doc = Document()
        for tema, texto in relatorios.items():
            doc.add_heading(tema, level=1)
            doc.add_paragraph(texto)

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        st.download_button(
            label="💾 Baixar Relatório em Word",
            data=buffer,
            file_name="relatorio_atividades.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    else:
        st.warning("Nenhuma coluna com conteúdo válido foi encontrada.")
